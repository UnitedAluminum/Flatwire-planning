"""
Render a client-facing markdown specification to a branded .docx.

    python DevelopmentPlan/Tools/build_docx.py <source.md> <output.docx> "<header title>"

With no arguments it rebuilds the PLC tag specification, which is the default below.

WHY THIS EXISTS
    The .md files in LatestDocument/RequirementDocuments/ are the source of truth; the
    .docx files in SRS/ are generated output. Edit the markdown and re-render - never
    edit the .docx. This script had to be re-derived three times because the pipeline
    was never committed, so it now lives here.

HOW THE BRANDING WORKS
    Nothing is recreated. SRS/PassScheduleGenerationSpec.docx is opened as a template
    and its body is stripped, keeping only the final <w:sectPr> - which is what carries
    the headerReference and footerReference. The output therefore inherits the header
    logo, the "United Aluminum Confidential" footer, the page-number fields, the page
    setup (8.5 x 11 at 0.6" margins) and the full style set, byte for byte.

WHAT IT HANDLES
    Headings (the first h1 becomes a centred title), tables with repeating shaded header
    rows, blockquote callouts, fenced code, ordered/unordered lists, "**Key:** value"
    metadata lines, and inline bold / italic / strike / code / links - the inline parser
    recurses, so **`nested code`** resolves rather than leaking its delimiters.
    A TOC field is inserted after the title; Word populates it on Update Field.

REQUIREMENTS
    python-docx (tested against 1.2.0). No pandoc, no Word automation.

GOTCHA
    If the output is open in Word the write fails with PermissionError - a ~$ lock file
    in SRS/ is the tell. Close it and re-run.
"""
import io, os, re, shutil, sys
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

ROOT = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..'))
TPL = os.path.join(ROOT, 'SRS', 'PassScheduleGenerationSpec.docx')

_args = sys.argv[1:]
SRC = os.path.abspath(_args[0]) if len(_args) > 0 else \
    os.path.join(ROOT, 'LatestDocument', 'RequirementDocuments', 'PLCTagSpecification.md')
OUT = os.path.abspath(_args[1]) if len(_args) > 1 else \
    os.path.join(ROOT, 'SRS', 'PLCTagSpecification.docx')
DOC_TITLE = _args[2] if len(_args) > 2 else \
    'Flat Wire - PLC Tag and Machine Interface Specification'
MONO = 'Consolas'
HDR_FILL = 'D9E2F3'      # table header row
QUOTE_FILL = 'FFF2CC'    # callout / blockquote
CODE_FILL = 'F2F2F2'


# --------------------------------------------------------------------------- helpers
def shade(el, hexfill):
    tcpr = el.get_or_add_tcPr() if el.tag.endswith('tc') else el
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hexfill)
    tcpr.append(sh)


def shade_para(p, hexfill):
    ppr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hexfill)
    ppr.append(sh)


def left_bar(p, color='2E74B5', sz=18):
    ppr = p._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    ln = OxmlElement('w:left')
    ln.set(qn('w:val'), 'single')
    ln.set(qn('w:sz'), str(sz))
    ln.set(qn('w:space'), '8')
    ln.set(qn('w:color'), color)
    bd.append(ln)
    ppr.append(bd)


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    trpr.append(el)


def keep_with_next(p):
    ppr = p._p.get_or_add_pPr()
    ppr.append(OxmlElement('w:keepNext'))


# ------------------------------------------------------------------ inline formatting
TOKEN = re.compile(
    r'(`[^`]+`)'                     # code
    r'|(\*\*\*.+?\*\*\*)'            # bold+italic
    r'|(\*\*.+?\*\*)'                # bold
    r'|(~~.+?~~)'                    # strike
    r'|(\*(?!\*).+?(?<!\*)\*)'       # italic
    r'|(\[[^\]]*\]\([^)]*\))'        # link
)


def add_runs(p, text, bold=False, italic=False, strike=False, link=False, _depth=0):
    """Emit runs for inline markdown.

    Recurses into bold/italic/strike spans so nested constructs such as
    **`PLC-Q##`** or ~~**FR-111**~~ resolve to a single correctly-styled run
    instead of leaking their delimiters.
    """
    if _depth == 0:
        text = text.replace('<br>', '\n').replace('&nbsp;', ' ')

    def plain(t):
        if not t:
            return
        r = p.add_run(t)
        r.bold, r.italic = bold, italic
        r.font.strike = strike
        if link:
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            r.font.underline = True

    pos = 0
    for m in TOKEN.finditer(text):
        if m.start() > pos:
            plain(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('`'):                       # code is a leaf - never recurse
            r = p.add_run(tok[1:-1])
            r.font.name = MONO
            r.font.size = Pt(9)
            r.bold = bold
            r.font.strike = strike
            r.font.color.rgb = RGBColor(0xA3, 0x1D, 0x1D)
        elif tok.startswith('***'):
            add_runs(p, tok[3:-3], True, True, strike, link, _depth + 1)
        elif tok.startswith('**'):
            add_runs(p, tok[2:-2], True, italic, strike, link, _depth + 1)
        elif tok.startswith('~~'):
            add_runs(p, tok[2:-2], bold, italic, True, link, _depth + 1)
        elif tok.startswith('*'):
            add_runs(p, tok[1:-1], bold, True, strike, link, _depth + 1)
        else:                                          # link -> label only
            lbl = re.match(r'\[([^\]]*)\]', tok).group(1)
            add_runs(p, lbl, bold, italic, strike, True, _depth + 1)
        pos = m.end()
    plain(text[pos:])


# ------------------------------------------------------------------------- TOC field
def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = r'TOC \o "1-3" \h \z \u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = 'Right-click and choose Update Field to build the table of contents.'
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for e in (f1, it, f2, t, f3):
        r._r.append(e)


# ----------------------------------------------------------------------- table build
def add_table(doc, rows, aligns):
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = doc.styles['Table Grid']
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for ri, cells in enumerate(rows):
        row = t.add_row()
        for ci in range(ncol):
            cell = row.cells[ci]
            cell.paragraphs[0].text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            txt = cells[ci] if ci < len(cells) else ''
            add_runs(p, txt, bold=(ri == 0))
            for r in p.runs:
                r.font.size = Pt(8.5)
            if ci < len(aligns) and aligns[ci] == 'c':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri == 0:
                shade(cell._tc, HDR_FILL)
        if ri == 0:
            repeat_header(row)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def split_row(ln):
    ln = ln.strip()
    if ln.startswith('|'):
        ln = ln[1:]
    if ln.endswith('|'):
        ln = ln[:-1]
    out, buf, i = [], '', 0
    while i < len(ln):                    # respect \| and inline code
        ch = ln[i]
        if ch == '\\' and i + 1 < len(ln) and ln[i + 1] == '|':
            buf += '|'; i += 2; continue
        if ch == '|':
            out.append(buf.strip()); buf = ''; i += 1; continue
        buf += ch; i += 1
    out.append(buf.strip())
    return out


def parse_aligns(ln):
    a = []
    for c in split_row(ln):
        c = c.strip()
        if c.startswith(':') and c.endswith(':'):
            a.append('c')
        elif c.endswith(':'):
            a.append('r')
        else:
            a.append('l')
    return a


# ============================================================================= main
def main():
    src = io.open(SRC, encoding='utf-8').read().replace('\r\n', '\n')
    lines = src.split('\n')

    shutil.copyfile(TPL, OUT)
    doc = docx.Document(OUT)

    # strip the template body, keeping the final sectPr (holds header/footer refs)
    body = doc.element.body
    sect = body.find(qn('w:sectPr'))
    for ch in list(body):
        if ch is not sect:
            body.remove(ch)
    print('template body stripped; sectPr kept')

    # header text
    hp = doc.sections[0].header.paragraphs[0]
    for r in hp.runs:
        if r.text.strip():
            r.text = DOC_TITLE
            break
    print('header set: ' + DOC_TITLE)

    i, n = 0, len(lines)
    first_h1 = True
    stats = {'h': 0, 'tbl': 0, 'quote': 0, 'code': 0, 'para': 0, 'list': 0}

    while i < n:
        ln = lines[i]

        # ---- fenced code
        if ln.startswith('```'):
            i += 1
            buf = []
            while i < n and not lines[i].startswith('```'):
                buf.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.12)
            shade_para(p, CODE_FILL)
            r = p.add_run('\n'.join(buf))
            r.font.name = MONO
            r.font.size = Pt(8)
            stats['code'] += 1
            continue

        # ---- table
        if ln.startswith('|') and i + 1 < n and re.match(r'^\|[\s:|-]+\|?\s*$', lines[i + 1]):
            head = split_row(ln)
            aligns = parse_aligns(lines[i + 1])
            i += 2
            rows = [head]
            while i < n and lines[i].startswith('|'):
                rows.append(split_row(lines[i])); i += 1
            add_table(doc, rows, aligns)
            stats['tbl'] += 1
            continue

        # ---- heading
        m = re.match(r'^(#{1,4})\s+(.*)$', ln)
        if m:
            lvl, txt = len(m.group(1)), m.group(2).strip()
            if first_h1 and lvl == 1:
                # title page
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(txt)
                r.bold = True
                r.font.size = Pt(20)
                r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)
                p.paragraph_format.space_after = Pt(14)
                first_h1 = False
                i += 1
                continue
            p = doc.add_paragraph(style='Heading %d' % min(lvl, 4))
            add_runs(p, txt)
            keep_with_next(p)
            stats['h'] += 1
            i += 1
            continue

        # ---- blockquote (group consecutive)
        if ln.startswith('>'):
            buf = []
            while i < n and (lines[i].startswith('>') or (buf and lines[i].strip() == '' and
                             i + 1 < n and lines[i + 1].startswith('>'))):
                buf.append(re.sub(r'^>\s?', '', lines[i]))
                i += 1
            first = True
            for b in buf:
                if not b.strip():
                    continue
                hm = re.match(r'^(#{1,4})\s+(.*)$', b)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.16)
                p.paragraph_format.space_before = Pt(6 if first else 2)
                p.paragraph_format.space_after = Pt(2)
                shade_para(p, QUOTE_FILL)
                left_bar(p)
                if hm:
                    add_runs(p, hm.group(2))
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(11)
                else:
                    add_runs(p, b)
                    for r in p.runs:
                        if not r.font.size:
                            r.font.size = Pt(9.5)
                first = False
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            stats['quote'] += 1
            continue

        # ---- horizontal rule -> ignore (headings already separate sections)
        if ln.strip() == '---':
            i += 1
            continue

        # ---- lists
        lm = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)$', ln)
        if lm:
            indent = len(lm.group(1))
            ordered = lm.group(2)[0].isdigit()
            style = 'List Number' if ordered else 'List Bullet'
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.28 + 0.22 * (indent // 2))
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, lm.group(3))
            for r in p.runs:
                if not r.font.size:
                    r.font.size = Pt(9.5)
            stats['list'] += 1
            i += 1
            continue

        # ---- blank
        if not ln.strip():
            i += 1
            continue

        # ---- label line: "**Key:** value" stands alone (the doc metadata block)
        LABEL = r'^\*\*[^*]+:\*\*'
        if re.match(LABEL, ln):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            add_runs(p, ln)
            for r in p.runs:
                if not r.font.size:
                    r.font.size = Pt(9.5)
            stats['para'] += 1
            i += 1
            continue

        # ---- paragraph (join soft-wrapped continuation lines)
        buf = [ln]
        i += 1
        while (i < n and lines[i].strip() and not lines[i].startswith(('|', '>', '#', '```'))
               and lines[i].strip() != '---'
               and not re.match(LABEL, lines[i])
               and not re.match(r'^(\s*)([-*]|\d+\.)\s+', lines[i])):
            buf.append(lines[i]); i += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        add_runs(p, ' '.join(buf))
        for r in p.runs:
            if not r.font.size:
                r.font.size = Pt(9.5)
        stats['para'] += 1

    # insert the TOC right after the title block (before the first Heading)
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading'):
            toc_p = doc.add_paragraph()
            r = toc_p.add_run('Contents')
            r.bold = True
            r.font.size = Pt(13)
            add_toc(doc)
            pb = doc.add_paragraph()
            pb.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
            # move the three new paragraphs ahead of the first heading
            newel = [doc.paragraphs[-3]._p, doc.paragraphs[-2]._p, doc.paragraphs[-1]._p]
            anchor = p._p
            for e in newel:
                e.getparent().remove(e)
                anchor.addprevious(e)
            break

    doc.save(OUT)
    print('stats:', stats)
    print('saved: ' + OUT)


main()
